import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class CVMakerApp:
    def __init__(self, root):
        self.root = root
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.root.title("Customizable CV Maker")
        self.root.geometry(f"{int(0.5*screen_width)}x{int(0.5*screen_height)}+{int(0.25*screen_width)}+{int(0.25*screen_height)}")

        # User data dictionary with max lengths
        self.user_data = {
            "name": tk.StringVar(),
            "email": tk.StringVar(),
            "phone": tk.StringVar(),
            "education": tk.StringVar(),
            "experience": tk.StringVar(),
            "skills": tk.StringVar(),
        }

        self.profile_pic_path = None  # To store the profile picture path

        # Create the form
        self.create_form()

    def create_form(self):
        ttk.Label(self.root, text="Customizable CV Maker", font=("Arial", 16)).pack(pady=10)

        # Name
        self.create_limited_entry("Name:", "name", 40)

        # Email
        self.create_limited_entry("Email:", "email", 40)

        # Phone
        self.create_limited_entry("Phone:", "phone", 20)

        # Education
        self.create_limited_entry("Education (separate with commas):", "education", 200)

        # Experience
        self.create_limited_entry("Work Experience (separate with commas):", "experience", 200)

        # Skills
        self.create_limited_entry("Skills (separate with commas):", "skills", 200)

        # Add profile picture
        ttk.Button(self.root, text="Add Profile Picture", command=self.add_profile_picture).pack(pady=10)
        self.profile_pic_label = ttk.Label(self.root, text="No picture selected.")
        self.profile_pic_label.pack()

        # Generate button
        ttk.Button(self.root, text="Generate CV", command=self.generate_cv).pack(pady=20)

    def create_limited_entry(self, label_text, field_name, max_length):
        """Creates a labeled entry with input length validation."""
        ttk.Label(self.root, text=label_text).pack(anchor="w", padx=20, pady=5)
        
        validate_command = (self.root.register(self.validate_length), "%P", max_length)
        ttk.Entry(
            self.root,
            textvariable=self.user_data[field_name],
            validate="key",
            validatecommand=validate_command,
        ).pack(fill="x", padx=20)

    @staticmethod
    def validate_length(new_value, max_length):
        """Validates that the input does not exceed the maximum length."""
        return len(new_value) <= int(max_length)

    def add_profile_picture(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Image files", "*.jpg *.jpeg *.png")],
            title="Select Profile Picture"
        )
        if file_path:
            self.profile_pic_path = file_path
            self.profile_pic_label.config(text=f"Selected: {file_path.split('/')[-1]}")

    def generate_cv(self):
        # Collect user data
        name = self.user_data["name"].get()
        email = self.user_data["email"].get()
        phone = self.user_data["phone"].get()
        education = self.user_data["education"].get().split(",")
        experience = self.user_data["experience"].get().split(",")
        skills = self.user_data["skills"].get().split(",")

        # Debugging Output
        print(f"Name: {name}, Email: {email}, Phone: {phone}")
        print(f"Education: {education}, Experience: {experience}, Skills: {skills}")

        if not name or not email or not phone:
            messagebox.showerror("Error", "Name, Email, and Phone are required!")
            return

        # Save DOCX
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            title="Save CV As"
        )
        if file_path:
            self.create_docx(file_path, name, email, phone, education, experience, skills)
            messagebox.showinfo("Success", f"CV saved as {file_path}")
        else:
            messagebox.showwarning("Warning", "No file path selected!")

    def create_docx(self, file_path, name, email, phone, education, experience, skills):
        doc = Document()

        # Add Title
        title = doc.add_heading("Curriculum Vitae", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add Name
        doc.add_paragraph(f"Name: {name.title()}", style="Heading 1")

        # Add Email
        doc.add_paragraph(f"Email: {email}", style="Normal")

        # Add Phone
        doc.add_paragraph(f"Phone: {phone}", style="Normal")

        # Add Education
        if education:
            doc.add_paragraph("Education:", style="Heading 2")
            for edu in education:
                doc.add_paragraph(f"- {edu.strip().title()}", style="Normal")

        # Add Experience
        if experience:
            doc.add_paragraph("Work Experience:", style="Heading 2")
            for exp in experience:
                doc.add_paragraph(f"- {exp.strip().title()}", style="Normal")

        # Add Skills
        if any(skill.strip() for skill in skills):
            doc.add_paragraph("Skills:", style="Heading 2")
            for skill in skills:
                doc.add_paragraph(f"- {skill.strip().title()}", style="Normal")

        # Save the DOCX
        doc.save(file_path)

# Create the Tkinter window and run the app
if __name__ == "__main__":
    root = tk.Tk()
    app = CVMakerApp(root)
    root.mainloop()
